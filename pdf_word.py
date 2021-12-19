#python處理PDF文件
#批量將PDF轉換為Word文檔

from itertools import filterfalse
import pdfplumber
from docx import Document
import os
from multiprocessing import Process


def convertPdf(fileName):
    with pdfplumber.open(fileName) as pdf:
        print("正在處理文件{0},一共{1}頁...".format(fileName,len(pdf.pages)))
        content = ''
        baseName = fileName.split('.')[0] 
        wordName = baseName + '.docx'
        flag = True
        if os.path.exists(wordName):
            os.remove(wordName)
        for i in range(len(pdf.pages)):
            print("正在處理<{0}>第{1}頁...".format(baseName,i))
            page = pdf.pages[i] #建立物件page，將pfd的頁數放進page
            if page.extract_text() == None:
                print("{0}是圖片拼接起來的，所以無法轉換")
                flag = False
                break
            #在('\n')後面加[:-1] 即截取從第0個到倒數第一個text
            page_content = '\n'.join(page.extract_text().split('\n')) 
            content = content + page_content
            if os.path.exists(wordName):
                doc = Document(wordName)
            else:
                doc = Document()
            doc.add_paragraph(content)
            doc.save(wordName)
            content=''
            print("<{0}>第{1}頁處理完成!".format(baseName,i))
        if flag:
            print("{0}處理完成,保存為{1}!!!!".format(fileName,wordName))


    
if __name__=='__main__':
    for file in os.listdir('.'):
        if os.path.isfile(file) and file.split('.')[1] == 'pdf': #file以'.'分開，如果'.'後面是'pdf'就進下一步
            p=Process(target=convertPdf, args=(file,)) #將process放進去
            p.start()

